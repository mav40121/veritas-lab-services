"""
VeritaPolicy Phase 2 template builder (data-driven).

Reads policy JSON files from server/policyTemplates/data/ and renders each
as a DOCX in the locked CFR-anchored format:
- Header: lab identity + policy metadata
- Section 1 Purpose (plain-CFR voice)
- Section 2 Regulatory Authority (eCFR text verbatim, citation + paragraph)
- Section 3 Scope
- Section 4 Policy (the lab's standing rule)
- Section 5 Procedure (operational steps)
- Section 6 Accreditor translation table (paraphrases for TJC/CAP/COLA/AABB)
- Section 7 Definitions
- Section 8 Revision history
- Page 1 director signature block (per CLAUDE.md §5)

Placeholder tokens (server replaces at download in Phase 3):
- <<LAB_NAME>>, <<CLIA_NUMBER>>, <<DIRECTOR_NAME>>, <<EFFECTIVE_DATE>>,
  <<RESPONSIBLE_ROLE>> (default: "Laboratory Director or Designee")

Run:
  python scripts/veritapolicy_phase2_template_builder.py            # all
  python scripts/veritapolicy_phase2_template_builder.py 11 20 5    # specific ids

Output:
  C:/Users/veril/OneDrive/Desktop/Lab/Verita Products/VeritaPolicy/templates/
"""

from __future__ import annotations
import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATA_DIR = Path(__file__).parent.parent / "server" / "policyTemplates" / "data"
OUTPUT_DIR = Path(r"C:\Users\veril\OneDrive\Desktop\Lab\Verita Products\VeritaPolicy\templates")

TEAL = RGBColor(0x01, 0x69, 0x6F)
TEAL_LIGHT = "E6F2F2"
INK = RGBColor(0x28, 0x25, 0x1D)
GREY = RGBColor(0x66, 0x66, 0x66)

def set_cell_shading(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def style_run(run, *, size=11, bold=False, italic=False, color=INK):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color

def add_para(doc, text="", *, size=11, bold=False, italic=False, color=INK, align=None, space_before=0, space_after=2):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        style_run(run, size=size, bold=bold, italic=italic, color=color)
    return p

def add_heading(doc, text, level=1):
    sizes = {1: 14, 2: 12, 3: 11}
    return add_para(doc, text, size=sizes.get(level, 11), bold=True, color=TEAL,
                    space_before=8 if level == 1 else 4, space_after=2)

def add_numbered(doc, items):
    for i, it in enumerate(items, 1):
        p = add_para(doc, "", size=10, space_after=2)
        run = p.add_run(f"  {i}.  ")
        style_run(run, size=10, bold=True)
        run = p.add_run(it)
        style_run(run, size=10)

def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = footer_para.add_run("VeritaPolicy™ | <<LAB_NAME>> | Confidential - For Internal Lab Use Only")
    style_run(run, size=8, color=GREY)
    tab = footer_para.add_run("\t\t")
    style_run(tab, size=8)
    run2 = footer_para.add_run("Page ")
    style_run(run2, size=8, color=GREY)
    fld_char1 = OxmlElement('w:fldChar'); fld_char1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = " PAGE "
    fld_char2 = OxmlElement('w:fldChar'); fld_char2.set(qn('w:fldCharType'), 'end')
    run3 = footer_para.add_run()
    run3._r.append(fld_char1); run3._r.append(instr); run3._r.append(fld_char2)
    style_run(run3, size=8, color=GREY)
    run4 = footer_para.add_run(" of "); style_run(run4, size=8, color=GREY)
    fld_char3 = OxmlElement('w:fldChar'); fld_char3.set(qn('w:fldCharType'), 'begin')
    instr2 = OxmlElement('w:instrText'); instr2.set(qn('xml:space'), 'preserve'); instr2.text = " NUMPAGES "
    fld_char4 = OxmlElement('w:fldChar'); fld_char4.set(qn('w:fldCharType'), 'end')
    run5 = footer_para.add_run()
    run5._r.append(fld_char3); run5._r.append(instr2); run5._r.append(fld_char4)
    style_run(run5, size=8, color=GREY)

def build_one(policy: dict, output_path: Path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    add_para(doc, "<<LAB_NAME>>", size=10, color=GREY, space_after=0)
    add_para(doc, "CLIA: <<CLIA_NUMBER>>", size=9, color=GREY, space_after=6)

    add_para(doc, policy["policy_name"].upper(), size=18, bold=True, color=TEAL, space_after=2)
    add_para(doc, f"Policy ID: VP-{policy['policy_id']}   |   Section: {policy.get('section','')}   |   Effective: <<EFFECTIVE_DATE>>",
             size=9, color=GREY, space_after=4)
    add_para(doc, "Approved by: <<DIRECTOR_NAME>>, Medical Director or Designee   |   Version: 1.0",
             size=9, color=GREY, space_after=10)

    # Page-1 signature
    sig_table = doc.add_table(rows=1, cols=3)
    sig_table.autofit = False
    widths = [Inches(2.7), Inches(2.2), Inches(1.3)]
    cells = sig_table.rows[0].cells
    for i, w in enumerate(widths):
        cells[i].width = w
    p = cells[0].paragraphs[0]
    r = p.add_run("Medical Director or Designee Signature"); style_run(r, size=8, bold=True, color=TEAL)
    p2 = cells[0].add_paragraph(); r2 = p2.add_run("________________________________________"); style_run(r2, size=10)
    p = cells[1].paragraphs[0]; r = p.add_run("Printed Name / Title"); style_run(r, size=8, bold=True, color=TEAL)
    p2 = cells[1].add_paragraph(); r2 = p2.add_run("<<DIRECTOR_NAME>>"); style_run(r2, size=10)
    p = cells[2].paragraphs[0]; r = p.add_run("Date"); style_run(r, size=8, bold=True, color=TEAL)
    p2 = cells[2].add_paragraph(); r2 = p2.add_run("____________"); style_run(r2, size=10)
    for c in cells:
        set_cell_shading(c, TEAL_LIGHT)

    add_para(doc, "", space_after=6)

    add_heading(doc, "1. PURPOSE", 1)
    add_para(doc, policy["purpose"], size=10)

    add_heading(doc, "2. REGULATORY AUTHORITY", 1)
    add_para(doc, "Verbatim text from the Code of Federal Regulations (eCFR). This section is the regulatory baseline; the lab does not modify it.",
             size=9, italic=True, color=GREY)
    for blk in policy.get("cfr_text_blocks", []):
        p = add_para(doc, "", size=10, space_before=4, space_after=2)
        r = p.add_run(f"{blk['citation']}"); style_run(r, size=10, bold=True, color=TEAL)
        r2 = p.add_run(f"  -  {blk.get('label','')}"); style_run(r2, size=10, bold=True, color=INK)
        body = doc.add_paragraph()
        body.paragraph_format.left_indent = Inches(0.3)
        body.paragraph_format.space_after = Pt(6)
        rb = body.add_run('"' + blk["verbatim"] + '"')
        style_run(rb, size=10, italic=True)

    add_heading(doc, "3. SCOPE", 1)
    add_para(doc, policy["scope"], size=10)

    add_heading(doc, "4. POLICY", 1)
    add_para(doc, "<<LAB_NAME>> adopts the following standing rule to satisfy the regulatory baseline in Section 2:",
             size=10, space_after=4)
    add_numbered(doc, policy["policy_statements"])

    add_heading(doc, "5. PROCEDURE", 1)
    add_para(doc, "Owner: <<RESPONSIBLE_ROLE>>", size=10, italic=True, color=GREY, space_after=4)
    add_numbered(doc, policy["procedure_steps"])

    # Section 6: Definitions (no accreditor crosswalk — CFR is the federal
    # baseline every CLIA lab must meet; accreditor-specific crosswalks are
    # the accreditor's published material, not safe for us to render.
    # The lab can map this CFR-anchored policy to their accreditor's
    # standards using the accreditor's own published crosswalk.)
    if policy.get("definitions"):
        add_heading(doc, "6. DEFINITIONS", 1)
        for term, defn in policy["definitions"]:
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
            r = p.add_run(f"{term}: "); style_run(r, size=10, bold=True)
            r2 = p.add_run(defn); style_run(r2, size=10)

    add_heading(doc, "7. REVISION HISTORY", 1)
    t = doc.add_table(rows=2, cols=4)
    t.style = "Light Grid Accent 1"
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(["Version", "Date", "Author", "Change"]):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]; r = p.add_run(h)
        style_run(r, size=9, bold=True, color=RGBColor(255, 255, 255))
        set_cell_shading(hdr_cells[i], "01696F")
    row1 = t.rows[1].cells
    for i, val in enumerate(["1.0", "<<EFFECTIVE_DATE>>", "<<DIRECTOR_NAME>>", "Initial adoption from VeritaPolicy generic template."]):
        row1[i].text = ""
        p = row1[i].paragraphs[0]; r = p.add_run(val); style_run(r, size=9)

    add_footer(doc)

    add_para(doc, "", space_after=6)
    add_para(doc, "VeritaPolicy™ generic templates are starting points. Final approval and clinical determination must be made by the laboratory director or designee. The laboratory remains responsible for tailoring this policy to its scope of services, accreditation status, and state requirements.",
             size=8, italic=True, color=GREY)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

def main():
    only_ids = set(sys.argv[1:]) if len(sys.argv) > 1 else None
    files = sorted(DATA_DIR.glob("*.json"))
    if not files:
        print(f"No JSON data files in {DATA_DIR}")
        return
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    built = 0
    for f in files:
        policy = json.loads(f.read_text(encoding="utf-8"))
        if only_ids and policy["policy_id"] not in only_ids:
            continue
        out = OUTPUT_DIR / f"VP-{policy['policy_id'].zfill(3)}_{policy['slug']}.docx"
        build_one(policy, out)
        built += 1
        print(f"  built  {out.name}")
    print(f"\nWrote {built} template(s) to {OUTPUT_DIR}")

if __name__ == "__main__":
    main()
