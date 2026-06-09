"""
Convert the San Carlos ABO/Rh forward/reverse typing assessment docx into
the JSON shape that VeritaComp's "Upload JSON" quiz dialog accepts.

Source: C:\\Users\\veril\\ABO_Rh_Typing_Assessment.docx
Output: C:\\Users\\veril\\ABO_Rh_Typing_Assessment_quiz.json

The output JSON is the format the existing
  NewQuizDialog "Upload JSON" handler reads:
    { meta: {title, ...}, questions: [{id, question, type, options, correct_answer, explanation}] }

Question prompts are HTML so the forward/reverse reaction table renders as
an actual <table>, not a comma-separated description. The HTML is sanitized
on the client side at render time (see DOMPurify import in VeritaCompAppPage).

Answer keys are derived programmatically from standard ABO/Rh interpretation
rules, with the derivation written into each question's `explanation` field
so the SME (Michael) can verify each one before uploading. Two questions
intentionally have a forward/reverse discrepancy and the correct answer is
"Further testing needed, send out to reference lab."

Run from anywhere:
    python scripts/convert_abo_rh_quiz.py

Requires python-docx installed in the active Python environment.
"""

from __future__ import annotations
import json
import os
from typing import List, Dict, Any
from docx import Document

SRC_PATH = r"C:\Users\veril\ABO_Rh_Typing_Assessment.docx"
OUT_PATH = r"C:\Users\veril\ABO_Rh_Typing_Assessment_quiz.json"


def interpret_reactions(anti_a: str, anti_b: str, anti_d: str, a1_cells: str, b_cells: str) -> Dict[str, Any]:
    """
    Apply standard ABO/Rh interpretation rules to the five reaction values.
    Returns the inferred blood type AND whether a forward/reverse discrepancy
    exists. Discrepancies should be answered "Further testing needed."

    Forward (cell) typing — what antigens are on the patient's red cells:
      Anti-A reactive (>=2+)  -> A antigen present
      Anti-B reactive (>=2+)  -> B antigen present
      Anti-D reactive (>=2+)  -> Rh+ (else Rh-)

    Reverse (serum) typing — what antibodies are in the patient's serum:
      A1 cells reactive (>=2+) -> anti-A present in serum
      B cells reactive (>=2+)  -> anti-B present in serum

    Concordance: serum antibodies must be against antigens the patient
    LACKS on their cells.
      O   serum has anti-A AND anti-B
      A   serum has anti-B only
      B   serum has anti-A only
      AB  serum has neither

    Any deviation = ABO discrepancy = refer out.
    """
    def reactive(g: str) -> bool:
        # 0 = nonreactive; "1+" through "4+" are reactive. Treat "2+" and
        # above as definitely reactive; 1+ is borderline (treat as reactive
        # here, but flag weak reactions as discrepancies if reverse
        # disagrees with forward).
        g = g.strip()
        if not g or g == "0":
            return False
        return True

    def weak(g: str) -> bool:
        g = g.strip()
        return g in {"1+", "2+"}

    a_pos = reactive(anti_a)
    b_pos = reactive(anti_b)
    rh_pos = reactive(anti_d)
    has_anti_a = reactive(a1_cells)
    has_anti_b = reactive(b_cells)

    # Forward type
    if a_pos and b_pos:
        forward = "AB"
    elif a_pos:
        forward = "A"
    elif b_pos:
        forward = "B"
    else:
        forward = "O"

    # Expected reverse for each forward type
    expected_reverse = {
        "O":  (True, True),    # has anti-A and anti-B
        "A":  (False, True),   # has anti-B only
        "B":  (True, False),   # has anti-A only
        "AB": (False, False),  # neither
    }[forward]

    actual_reverse = (has_anti_a, has_anti_b)
    concordant = actual_reverse == expected_reverse

    # Also count weak reverse reactions in an "AB" forward as a discrepancy
    # (a 2+ anti-A1 in an apparent AB patient = A2B subtype with anti-A1;
    # at the tech level the correct answer is "Further testing").
    if forward == "AB" and (weak(a1_cells) or weak(b_cells)):
        concordant = False

    rh_label = "Positive" if rh_pos else "Negative"
    inferred_type = f"{forward} {rh_label}"

    return {
        "forward": forward,
        "rh_positive": rh_pos,
        "expected_reverse": expected_reverse,
        "actual_reverse": actual_reverse,
        "concordant": concordant,
        "inferred_type": inferred_type,
    }


def reaction_table_html(anti_a: str, anti_b: str, anti_d: str, a1_cells: str, b_cells: str) -> str:
    """
    Render the reaction table as HTML so the quiz player can show it
    inline with the question prompt. DOMPurify on the client whitelists
    <table>/<thead>/<tbody>/<tr>/<th>/<td>.
    """
    return (
        '<table class="reaction-table"><thead>'
        '<tr><th colspan="3">Forward (Cell) Typing</th><th colspan="2">Reverse (Serum) Typing</th></tr>'
        '<tr><th>Anti-A</th><th>Anti-B</th><th>Anti-D</th><th>A1 Cells</th><th>B Cells</th></tr>'
        '</thead><tbody><tr>'
        f'<td>{anti_a}</td><td>{anti_b}</td><td>{anti_d}</td><td>{a1_cells}</td><td>{b_cells}</td>'
        '</tr></tbody></table>'
    )


def walk_docx(path: str):
    doc = Document(path)
    body = doc.element.body
    para_iter = iter(doc.paragraphs)
    table_iter = iter(doc.tables)
    for child in body.iterchildren():
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            p = next(para_iter)
            t = p.text.strip()
            if t:
                yield ("P", t, None)
        elif tag == 'tbl':
            tbl = next(table_iter)
            rows = [[c.text.strip() for c in row.cells] for row in tbl.rows]
            yield ("T", None, rows)


def main():
    if not os.path.exists(SRC_PATH):
        raise SystemExit(f"Source docx not found: {SRC_PATH}")

    items: List[Any] = list(walk_docx(SRC_PATH))

    questions: List[Dict[str, Any]] = []
    title: str = ""
    subtitle: str = ""

    # Walk pattern: title paragraph(s), then for each question:
    #   P "Patient N"
    #   T reaction table
    #   P "A.  Further testing..."  ... P "F.  ..."
    i = 0
    while i < len(items):
        kind, text, table = items[i]
        if kind == "P" and text and not text.startswith("Patient "):
            if not title:
                title = text
            elif not subtitle:
                subtitle = text
            else:
                # Tertiary intro paragraph — append to subtitle.
                subtitle += " | " + text
            i += 1
            continue

        if kind == "P" and text and text.startswith("Patient "):
            patient_label = text
            # Next item should be the reaction table.
            i += 1
            if i >= len(items) or items[i][0] != "T":
                raise SystemExit(f"Expected reaction table after '{patient_label}', got {items[i]}")
            _, _, table = items[i]
            i += 1
            # Reaction values live in the third row, cols 0..4 (col 5 is the
            # spacer cell created by the merged header layout).
            reactions = table[2][:5]
            anti_a, anti_b, anti_d, a1_cells, b_cells = reactions
            # Next 6 paragraphs are A-F choices.
            options: List[str] = []
            for _ in range(6):
                if i >= len(items) or items[i][0] != "P":
                    break
                _, opt_text, _ = items[i]
                if opt_text and opt_text[:2] in {f"{ch}." for ch in "ABCDEF"}:
                    # Strip the leading "A." / "B." marker; the player adds
                    # its own labels.
                    options.append(opt_text.split(".", 1)[1].strip())
                    i += 1
                else:
                    break
            if len(options) != 6:
                raise SystemExit(f"Expected 6 options for {patient_label}, got {len(options)}: {options}")

            # Apply interpretation
            interp = interpret_reactions(anti_a, anti_b, anti_d, a1_cells, b_cells)
            # Find the correct option letter:
            #  - If discrepancy, answer is A ("Further testing needed...")
            #  - Else find the option whose text matches inferred_type.
            if not interp["concordant"]:
                correct_letter = "A"
                explanation = (
                    f"Forward type is {interp['forward']} and Rh "
                    f"{'positive' if interp['rh_positive'] else 'negative'}, "
                    "but the reverse type does not match the expected pattern "
                    "for that forward type. ABO discrepancy: refer to reference lab."
                )
            else:
                inferred = interp["inferred_type"]
                correct_letter = None
                for idx, opt_text in enumerate(options):
                    if opt_text.strip().lower() == inferred.lower():
                        correct_letter = "ABCDEF"[idx]
                        break
                if correct_letter is None:
                    raise SystemExit(
                        f"For {patient_label}, inferred type '{inferred}' "
                        f"not found in options: {options}"
                    )
                explanation = (
                    f"Forward type {interp['forward']}, Rh "
                    f"{'positive' if interp['rh_positive'] else 'negative'}; "
                    "reverse type matches expected pattern. "
                    f"Concordant interpretation: {inferred}."
                )

            qid = f"q{len(questions) + 1}"
            prompt_html = (
                f'<p>Interpret the ABO group and Rh type for <strong>{patient_label}</strong>.</p>'
                + reaction_table_html(anti_a, anti_b, anti_d, a1_cells, b_cells)
            )
            questions.append({
                "id": qid,
                "question": prompt_html,
                "type": "multiple_choice",
                "options": [f"{ch}. {opt}" for ch, opt in zip("ABCDEF", options)],
                "correct_answer": correct_letter,
                "explanation": explanation,
            })
            continue

        i += 1

    out = {
        "meta": {
            "title": title or "ABO/Rh Forward and Reverse Typing Assessment",
            "subtitle": subtitle,
            "source": "ABO_Rh_Typing_Assessment.docx",
            "question_format": "html",
            "draft_answer_keys": True,
            "verifier_note": (
                "Answer keys were derived programmatically from standard ABO/Rh "
                "interpretation rules. Each question's `explanation` field shows "
                "the derivation. Verify all 10 before uploading."
            ),
        },
        "questions": questions,
    }

    with open(OUT_PATH, "w", encoding="utf-8") as fh:
        json.dump(out, fh, indent=2, ensure_ascii=False)

    print(f"Wrote {len(questions)} questions to {OUT_PATH}")
    for q in questions:
        # Strip the HTML prompt down to the patient label for the console
        # summary; keep the options + answer key visible.
        label = q["question"].split("<strong>")[1].split("</strong>")[0] if "<strong>" in q["question"] else q["id"]
        print(f"  {label:<11} -> {q['correct_answer']}: {q['options'][ord(q['correct_answer']) - ord('A')]}")


if __name__ == "__main__":
    main()
